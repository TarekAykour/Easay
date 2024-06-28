import docx
from pyrtf.ng import Document, Section, Paragraph, Renderer
import os

def docx_to_rtf(docx_file, rtf_file):
    # Load the .docx file
    doc = docx.Document(docx_file)
    
    # Create an RTF document
    rtf_doc = Document()
    section = Section()
    rtf_doc.Sections.append(section)
    
    # Convert paragraphs from .docx to .rtf
    for para in doc.paragraphs:
        rtf_para = Paragraph(para.text)
        section.append(rtf_para)
    
    # Save the RTF document
    renderer = Renderer()
    with open(rtf_file, 'w', encoding='utf-8') as f:
        renderer.Write(rtf_doc, f)


