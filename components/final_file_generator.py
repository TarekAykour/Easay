from components.file_utils import is_image, read_doc, read_rtf
import re
from docx import Document
from components.APA import url_to_apa

def combine_files(files):
    combined_file_path = './static/final_file/Final_essay.doc'  # Changed extension to .docx
    titles = []
    urls = []
    all_content = []
    image_filenames = []

    # Create a new Document
    doc = Document()

    for file_path in files:
        if is_image(file_path):
            all_content.append(f'[IMAGE: {file_path}]')
        # for .rtf files
        elif file_path.endswith('.rtf'):
            title, paragraphs, file_urls, image_filenames = read_rtf(file_path)
            # image_filenames.extend(file_images)
            if title:
                titles.append(title)
                all_content.append(("Subtitle", title))
            for paragraph in paragraphs:
                if re.search(r'\\fs([2-9]\d|[1-9]\d\d)', paragraph):  # Detect font size 24 or greater
                    all_content.append(("Subtitle", paragraph))
                    
                    
                if re.match(r'^\d+\.', paragraph):
                    doc.add_paragraph("\n")
                    doc.add_paragraph(paragraph)
                else:
                    all_content.append(("Content", paragraph))
                    
            urls.extend(file_urls)
        # for .docx and .doc files
        elif file_path.endswith('.docx') or file_path.endswith('.doc'):
            title, paragraphs, file_urls = read_doc(file_path)
            if title:
                titles.append(title)
                all_content.append(("Subtitle", title))
            for paragraph in paragraphs:
                all_content.append(("Content", paragraph))
            urls.extend(file_urls)

    # Write the table of contents
    doc.add_heading("Table of Contents", level=1)
    for i, title in enumerate(titles, 1):
        doc.add_heading(f"{i}. {title}", level=2)
    
    # Write the content
    for content_type, content in all_content:
        if content_type == "Subtitle":
            doc.add_heading(content, level=1)
        elif content_type == "Content":
            if re.match(r'^\d+\.', content):
                doc.add_paragraph("\n" + content)
            else:
                doc.add_paragraph(content)

        # Add images to the document
        if '[IMAGE:' in content:
            image_filename = content.replace('[IMAGE: ', '').replace(']', '')
            if image_filename in image_filenames:
                doc.add_picture(image_filename)

    # Write the URLs
    doc.add_heading("URL Section", level=1)
    for url in urls:
        doc.add_paragraph(url_to_apa(url))

    # Save the document
    doc.save(combined_file_path)
    print(f"Combined file saved to {combined_file_path}")
